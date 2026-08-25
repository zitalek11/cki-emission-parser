from __future__ import annotations

from pathlib import Path

import fitz
import pytest
from docx import Document


@pytest.fixture
def digital_pdf(tmp_path: Path) -> Path:
    path = tmp_path / "issuance_decision.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_htmlbox(
        page.rect,
        "<p>РЕШЕНИЕ О ВЫПУСКЕ ЦЕННЫХ БУМАГ</p>"
        "<p>Полное фирменное наименование эмитента: Публичное акционерное общество «Пример»</p>"
        "<p>Номинальная стоимость каждой ценной бумаги: 1 000 (одна тысяча) российских рублей.</p>",
    )
    doc.save(path)
    doc.close()
    return path


@pytest.fixture
def scan_like_pdf(tmp_path: Path) -> Path:
    path = tmp_path / "scan.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.draw_rect(page.rect, color=(0.8, 0.8, 0.8), fill=(0.9, 0.9, 0.9))
    doc.save(path)
    doc.close()
    return path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    path = tmp_path / "placement_terms.docx"
    document = Document()
    document.add_paragraph("Документ, содержащий условия размещения")
    document.add_paragraph("Дата начала размещения: 1 января 2026 г.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Параметр"
    table.cell(0, 1).text = "Значение"
    table.cell(1, 0).text = "Количество"
    table.cell(1, 1).text = "10 000"
    document.save(path)
    return path
